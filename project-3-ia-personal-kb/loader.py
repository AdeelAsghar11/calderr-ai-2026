"""
loader.py — Multi-format Document Loader

Reads every supported file from docs/ and returns LangChain Document
objects, one per file, with metadata attached (source name, file type,
word count).

Supported formats:
  .pdf   — extracted page by page with pypdf
  .md    — read directly (already clean text)
  .txt   — read directly
  .docx  — extracted paragraph by paragraph with python-docx

Any other file extension is silently skipped — this lets you keep
screenshots, zips, or other files in the same folder without breaking
ingestion.

Each Document's metadata includes:
  source     — filename without extension (used for --source filtering)
  filename   — full filename with extension (shown in citations)
  file_type  — pdf / md / txt / docx
  word_count — total words in the document before chunking
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document


def load_pdf(path: Path) -> str:
    """Extract text from a PDF, page by page, using pypdf (pure Python)."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages  = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n\n".join(pages)


def load_docx(path: Path) -> str:
    """Extract text from a Word document, paragraph by paragraph."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def load_text(path: Path) -> str:
    """Read a plain text or markdown file directly."""
    return path.read_text(encoding="utf-8", errors="ignore")


LOADERS = {
    ".pdf":  load_pdf,
    ".docx": load_docx,
    ".md":   load_text,
    ".txt":  load_text,
}


def load_documents(docs_dir: str = "./docs") -> list[Document]:
    """
    Load every supported file in docs_dir into a list of Document objects.

    Files with unsupported extensions are skipped silently.
    Files that fail to extract (corrupted PDF, empty file) are skipped
    with a printed warning rather than crashing the whole ingestion run.
    """
    docs_path = Path(docs_dir)
    if not docs_path.exists():
        raise FileNotFoundError(f"Docs directory not found: {docs_dir}")

    documents = []

    for path in sorted(docs_path.iterdir()):
        if not path.is_file():
            continue

        ext    = path.suffix.lower()
        loader = LOADERS.get(ext)
        if loader is None:
            continue  # unsupported file type — skip quietly

        try:
            text = loader(path)
        except Exception as e:
            print(f"  [skip] {path.name}: extraction failed ({e})")
            continue

        if not text.strip():
            print(f"  [skip] {path.name}: no extractable text")
            continue

        documents.append(Document(
            page_content=text,
            metadata={
                "source":     path.stem.lower().replace(" ", "_").replace("-", "_"),
                "filename":   path.name,
                "file_type":  ext.lstrip("."),
                "word_count": len(text.split()),
            },
        ))

    return documents


# ── Quick test ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    docs = load_documents("./docs")
    print(f"\nLoaded {len(docs)} documents:\n")
    for d in docs:
        print(f"  {d.metadata['filename']:<45} {d.metadata['file_type']:>5}  {d.metadata['word_count']:>6} words")
